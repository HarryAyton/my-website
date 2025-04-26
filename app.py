from flask import Flask, render_template, request, send_file
from docx import Document
import os

app = Flask(__name__)

# Sample questions
questions = {
    "Maths": {
        "Vectors": [
            "Prove that vectors a, b, and c are coplanar.",
            "Find the angle between two vectors."
        ],
        "Calculus": [
            "Differentiate f(x) = x^2 * sin(x)."
        ]
    },
    "Legal Studies": {
        "Human Rights": [
            "Explain the role of the United Nations in enforcing human rights."
        ]
    }
}

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        subject = request.form['subject']
        topic = request.form['topic']
        doc = Document()
        doc.add_heading(f"{subject} - {topic}", level=1)

        for q in questions[subject][topic]:
            doc.add_paragraph(q, style='List Bullet')

        filename = f"{subject}_{topic}_questions.docx"
        filepath = os.path.join("docs", filename)
        doc.save(filepath)

        return render_template('success.html', filename=filename)

    return render_template('index.html', questions=questions)

@app.route('/download/<filename>')
def download(filename):
    return send_file(os.path.join("docs", filename), as_attachment=True)

if __name__ == '__main__':
    os.makedirs("docs", exist_ok=True)
    app.run(debug=True)
